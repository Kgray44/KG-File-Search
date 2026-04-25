"""DOCX extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from kgfs.extractors.base import ExtractionResult, failed, ok


def extract_docx(path: Path) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError:
        return failed("python-docx is not installed")

    try:
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return ok("\n".join(part for part in parts if part))
    except Exception as exc:
        return failed(str(exc))

